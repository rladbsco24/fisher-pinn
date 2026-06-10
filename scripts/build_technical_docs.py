from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
BLACK = RGBColor(0, 0, 0)
FONT = "Malgun Gothic"
BODY_SIZE = Pt(10)


def _set_rfonts(obj, font_name: str = FONT) -> None:
    rpr = obj._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _format_run(run, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = BLACK
    run.font.bold = bold
    run.font.italic = italic
    _set_rfonts(run)


def _set_paragraph_style(paragraph, *, before: float = 0.0, after: float = 4.0, line_spacing: float = 1.05) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def _configure_document(document: Document, title: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = BODY_SIZE
        style.font.color.rgb = BLACK
        _set_rfonts(style)
        style.paragraph_format.line_spacing = 1.05
    styles["Normal"].paragraph_format.space_after = Pt(4)
    for heading in ("Heading 1", "Heading 2", "Heading 3"):
        styles[heading].font.bold = True
        styles[heading].paragraph_format.space_before = Pt(4)
        styles[heading].paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_style(header, after=0)
    for run in header.runs:
        _format_run(run, bold=False)


def _add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_style(paragraph, after=2)
    run = paragraph.add_run(title)
    _format_run(run, bold=True)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_style(sub, after=8)
    run = sub.add_run(subtitle)
    _format_run(run)


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    _set_paragraph_style(paragraph, before=4, after=4)
    run = paragraph.add_run(text)
    _format_run(run, bold=True)


def _add_para(document: Document, text: str, *, after: float = 4.0) -> None:
    paragraph = document.add_paragraph()
    _set_paragraph_style(paragraph, after=after)
    run = paragraph.add_run(text)
    _format_run(run)


def _set_cell_text(cell, text: str, *, bold: bool = False, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    _set_paragraph_style(paragraph, before=0, after=0, line_spacing=1.0)
    paragraph.text = ""
    run = paragraph.add_run(text)
    _format_run(run, bold=bold)


def _set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for tag, value in (
        ("top", "90"),
        ("bottom", "90"),
        ("left", "140"),
        ("right", "140"),
        ("start", "140"),
        ("end", "140"),
    ):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[float]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    total_dxa = int(round(sum(widths) * 1440))
    width_dxa = [int(round(width * 1440)) for width in widths]

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in width_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width_dxa[idx]))


def _add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    _set_cell_margins(table)
    _set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        table.columns[idx].width = Inches(widths[idx])
        _set_cell_text(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            _set_cell_text(cells[idx], value)
    _set_table_geometry(table, widths)
    spacer = document.add_paragraph()
    _set_paragraph_style(spacer, after=4)


def _add_reference(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    _set_paragraph_style(paragraph, before=0, after=3, line_spacing=1.0)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    run = paragraph.add_run(text)
    _format_run(run)


PINN_PAGES: list[tuple[str, list[str], tuple[list[str], list[list[str]], list[float]] | None]] = [
    (
        "1. 구현 범위와 핵심 구조",
        [
            "현재 저장소는 기본 Fisher-KPP PINN과 한국 산림청 pine-wilt PINN을 하나의 반응-확산 모델 계열로 구성한다. 모델은 Fisher-KPP 방정식, known initial condition, moving-front 보정 손실, 계수 추정, RK4 기준해 비교를 통합해 forward field와 front geometry를 함께 학습한다.",
            "기본 계열은 정규화된 2차원 Fisher-KPP forward 문제를 대상으로 한다. 코드상 기본 실행은 `scripts/run_inverse_origin.py`의 `geo_spectral_forward` 프로파일이며, 이전 inverse-origin 데모보다 forward 해석, known initial condition, front-aware loss, adaptive sampling, RK4 비교를 더 명시적으로 포함한다.",
            "산림청 계열은 같은 반응-확산 구조를 한국 pine-wilt 관측 밀도 격자에 적용한다. 구현은 land mask, sea exclusion, 조치 전 월별 시간축, 물리 계수 변환, 관측 밀도 가중치를 동일한 forward PINN/RK4 비교 체계 안에서 처리한다.",
        ],
        (
            ["항목", "현재 구현 기준"],
            [
                ["기본문제", "Fisher-KPP reaction-diffusion, 필요 시 advection을 제거한 forward PINN"],
                ["기본 PINN", "PirateNet/RWF 스타일 backbone, Fourier/front feature, hard IC, front-aware losses"],
                ["산림청 PINN", "land-only collocation, sea penalty, D/r 학습, 공간 계수장, 월별 조치 전 비교"],
            ],
            [1.45, 5.60],
        ),
    ),
    (
        "2. 기본 Fisher-KPP 문제와 기준해",
        [
            "기본 PDE는 u_t = D Laplacian(u) + r u(1-u)이다. Gaussian seed 문제에서는 정규화된 단위 정사각형에서 초기 국소 감염원이 퍼지는 형태를 사용하고, Ablowitz-Zeppetella 계열 검증에서는 u(x,t) = (1 + exp((x - c t - x0)/sqrt(6)))^-2, c = 5/sqrt(6)인 알려진 travelling wave를 기준으로 삼는다.",
            "정규화는 코드가 공간을 [0, 1] 좌표로 다루기 때문에 중요하다. 예를 들어 AZ 1D 기준은 x in [-20, 20]을 단위 좌표로 변환하면서 D_norm = 1 / 40^2로 들어간다. 이 때문에 수치적으로 보이는 diffusion 값은 물리식의 D=1과 다르지만, 변환을 거친 PDE는 같은 front speed와 wave profile을 나타낸다.",
            "기본 PINN과 RK4의 오차 차이는 두 방법의 계산 대상에서 발생한다. RK4는 알려진 초기장과 PDE를 고정 격자에서 직접 시간 적분하고, PINN은 연속 좌표 함수 u(x,y,t)를 신경망으로 근사하면서 PDE residual, 관측, 초기조건, front geometry를 동시에 만족시킨다. RK4는 forward reference solver로 높은 격자 정확도를 제공하고, PINN은 inverse/irregular data/연속 surrogate 기능을 제공한다.",
        ],
        (
            ["기준", "의미"],
            [
                ["AZ exact wave", "front phase와 low-level ring을 검증하기 위한 정확한 travelling-wave 기준"],
                ["Gaussian seed", "국소 발병원이 2D에서 확산되는 synthetic forward benchmark"],
                ["RK4 reference", "PINN 최종장과 front metrics를 비교하는 같은 문제의 수치 기준"],
            ],
            [1.45, 5.60],
        ),
    ),
    (
        "3. 기본 PINN의 모델 구조",
        [
            "기본 forward PINN은 `OriginPINN`을 사용하며 `geo_spectral_forward` 프로파일에서 architecture는 pirate, Fourier sigma는 완화된 값, random weight factorization은 켜진다. 입력은 x, y, t와 함께 seed-relative geometry, travelling-wave 좌표, front-local Fourier feature, KPP front envelope, 공간 계수장을 사용한다.",
            "hard initial condition은 t=0 부근에서 네트워크 출력이 초기장을 쉽게 무시하지 못하도록 한다. 이전 결과에서 PINN이 처음부터 퍼진 haze처럼 보였던 현상은 sparse front 문제에서 MSE와 PDE residual만으로는 대부분의 영역을 낮은 값으로 흐리게 만드는 해가 벌점을 적게 받을 수 있기 때문에 발생했다. 현재 구현은 known IC loss와 hard IC를 동시에 써서 이 현상을 줄인다.",
            "D와 r은 기본 forward 프로파일에서 학습 가능하다. 학습 계수는 physics parameter anchor, 공간 계수 regularization, front-speed 관련 loss로 안정화된다. 이 설계는 산림청 zip에서 D_tilde와 r_tilde를 PINN으로 추정한 뒤 D_phys = D_tilde * S_km^2 / DT, r_phys = r_tilde / DT로 해석하던 원리를 확장한다.",
        ],
        (
            ["구성", "역할"],
            [
                ["Pirate/RWF backbone", "front와 residual이 강한 비선형장에서도 gradient 흐름을 안정화"],
                ["front-local features", "예상 front 주변의 moving coordinate를 네트워크 입력에 추가"],
                ["spatial coefficients", "D(x,y), r(x,y)의 약한 보정장을 허용하되 regularization으로 제한"],
            ],
            [1.65, 5.40],
        ),
    ),
    (
        "4. front-aware 손실과 moving-front 보정",
        [
            "front-level-set alignment loss는 expected low-level ring 주변에서 u의 level, 안쪽/바깥쪽 순서 관계, front-normal slope를 함께 제약한다. Fisher-KPP leading edge가 만드는 등치선 위치를 soft geometric regularizer로 사용해 front phase와 front thickness를 직접 정렬한다.",
            "front profile loss는 expected front corridor에서 선형화된 KPP leading-edge profile을 맞추도록 한다. low-level area loss만 쓰면 넓은 저농도 haze가 통과할 수 있으므로, support Tversky, contrast, mass floor, mass balance를 함께 둔다. 이는 대부분을 0으로 예측하는 해와 넓게 흐린 해를 모두 비용이 큰 방향으로 밀어내기 위한 장치다.",
            "학습은 time marching과 time slab curriculum을 포함한다. 초반에는 쉬운 시간 구간과 초기조건을 먼저 고정하고, 뒤로 갈수록 전체 시간축, front loss, PDE residual을 확대한다. RAR와 front-aware sampler는 residual, u(1-u), gradient 정보를 결합해 active front 주변 collocation을 보강한다.",
        ],
        (
            ["손실/샘플링", "보정 대상"],
            [
                ["level-set/profile", "front phase가 늦거나 빠지고 halo가 생기는 문제"],
                ["support/contrast/mass", "all-zero 해 또는 넓은 저농도 haze 해"],
                ["time marching/RAR", "parabolic PDE의 장시간 누적 오차와 front-local residual 누락"],
            ],
            [1.75, 5.30],
        ),
    ),
    (
        "5. 기본 PINN 검증, 주요 문제, 해결 방안",
        [
            "현재 코드의 full run은 기본적으로 1200 epochs이며, quick run은 60 또는 120 epochs로 빠른 회귀 검증을 수행한다. 최근 수정 전 full run은 D = 0.01959, r = 3.01315처럼 물리 계수는 목표값에 가깝게 맞췄지만 final-time relative L2 = 0.4135로 RK4 reference의 0.00147과는 큰 차이를 남겼다. 이 결과는 PINN이 계수와 질량을 맞추더라도 moving front의 phase와 final field shape를 별도로 제어해야 함을 보여준다.",
            "가장 최근에 확인된 직접적인 학습 절차 문제는 best checkpoint 선택 기준이었다. 기존 구현은 validation observation MSE만 기준으로 모델을 복원했기 때문에, 1200 epochs를 학습하고도 epoch 1의 낮은 관측점 MSE checkpoint가 최종 모델로 선택될 수 있었다. 이 문제는 validation MSE, late-time physics proxy, PDE residual, known initial condition, front proxy, mass proxy를 함께 보는 composite checkpoint score와 최소 checkpoint epoch 조건으로 수정했다.",
            "주요 관찰은 front area와 mass metric이 field L2보다 front failure를 더 선명하게 드러낸다는 점이다. error map은 active front 주변에서 가장 높은 정보를 제공하며, support, contrast, mass floor, IC는 all-zero 해와 diffuse haze 해를 억제한다. D/r 학습은 forward field와 front speed의 물리 일관성을 함께 추정하는 핵심 구성이지만, Fisher-KPP front speed가 주로 2 sqrt(D r)에 의해 정해지므로 계수 식별성과 field profile 정렬은 별도 검증으로 관리한다.",
        ],
        (
            ["주요 문제", "원인", "해결 방안"],
            [
                ["초기 checkpoint 복원", "관측점 MSE만 낮은 초기 모델이 최종장, PDE, front보다 우선됨", "composite checkpoint score와 최소 epoch 조건 적용"],
                ["diffuse haze", "sparse front에서 낮은 농도를 넓게 깔아도 MSE 벌점이 작음", "known IC, hard IC, front profile, contrast, mass floor 강화"],
                ["all-zero collapse", "대부분의 영역이 0인 해가 관측 sparse regime에서 유리해질 수 있음", "support Tversky, mass trajectory, density-weighted data loss 적용"],
                ["front phase error", "moving front가 조금만 어긋나도 relative L2가 크게 증가", "level-set alignment, front-normal profile, time marching 적용"],
                ["twin-trap error band", "front 안쪽과 바깥쪽 순서 관계가 약하고 halo가 남음", "inside/outside hinge, normal-slope, front-aware sampling 적용"],
                ["D/r 식별성", "front speed는 D와 r의 곱에 민감하고 개별 계수는 약하게 식별됨", "physics anchor, coefficient regularization, learned physics reporting 적용"],
                ["RK4와 큰 정확도 차이", "PINN은 연속 surrogate 최적화이고 RK4는 고정 격자 forward 적분", "RK4 reference, discrete RK4 consistency, final-field/front/mass 공동 검증 적용"],
            ],
            [1.45, 2.55, 3.25],
        ),
    ),
    (
        "6. 실패 모드별 진단 근거",
        [
            "기본 PINN에서 발생한 문제는 단일 loss 값만으로 판단하지 않는다. 관측점 MSE, PDE residual, final-time field error, front-area metric, active-front band, mass trajectory, learned D/r, error map, GIF를 함께 보고 어느 실패 모드가 우세한지 분류한다. 이 방식은 PINN이 관측점에서는 좋아 보이지만 전체장에서는 틀리는 상황을 분리하기 위한 절차다.",
            "초기 checkpoint 복원 문제는 validation observation MSE가 가장 낮은 시점과 실제 final-field 성능이 일치하지 않는 데서 드러났다. 1200 epochs를 실행했는데도 epoch 1이 복원된 것은 최적화가 실패했다는 뜻이 아니라, 모델 선택 기준이 학습 목적과 어긋났다는 의미다. 따라서 현재 보고 기준은 checkpoint epoch, checkpoint score, validation MSE, PDE residual, teacher proxy를 함께 기록한다.",
            "front phase error와 twin-trap error band는 absolute-error map에서 active front 주변에 쌍봉형 또는 고리형 오차가 생기는 형태로 나타난다. 이 현상은 front의 면적만 맞추고 안쪽-바깥쪽 순서, level value, normal slope를 동시에 맞추지 못할 때 발생한다. field L2가 커지는 주된 원인이므로 front contour, active-front area, low-level support recall을 별도로 확인한다.",
            "산림청 PINN에서는 mean relative L2와 mass error만으로 충분하지 않다. 최근 full run에서 PINN은 mass error를 줄였지만 support false-negative가 남아 실제 감염 support를 보수적으로 놓칠 수 있었다. 따라서 산림청 결과는 relative L2, correlation, mass error, support FNR/FPR, Dice/Tversky, sea leakage 여부를 함께 제시한다.",
        ],
        (
            ["진단 항목", "문제가 드러나는 신호", "확인 산출물"],
            [
                ["checkpoint mismatch", "best epoch가 지나치게 이르고 final field가 좋지 않음", "metrics.json, training_diagnostics.png"],
                ["diffuse haze", "넓은 저농도 영역이 생기며 front contour가 흐려짐", "reconstruction.png, spacetime_error.png"],
                ["all-zero collapse", "관측 MSE는 작지만 mass와 support가 급감", "mass trajectory, support metric"],
                ["front phase error", "front 앞뒤로 큰 absolute-error band가 형성", "pinn_vs_rk4_comparison.png"],
                ["twin-trap band", "front 양쪽에 대칭적 고리형 error가 남음", "residual_front_diagnostics.png"],
                ["D/r ambiguity", "front speed는 맞지만 D와 r의 개별값이 흔들림", "learned physics, coefficient stats"],
                ["Korea support miss", "L2/mass는 개선되나 FNR이 높음", "korea_error_baselines.gif, support metrics"],
            ],
            [1.65, 3.15, 2.45],
        ),
    ),
    (
        "7. 해결 방안의 구현 반영",
        [
            "해결 방안은 loss를 단순히 늘리는 방식이 아니라, 실패 모드와 직접 연결되는 제약을 추가하는 방식으로 구성했다. known initial condition과 hard IC는 초기장 붕괴를 막고, level-set/profile/support loss는 front 위치와 두께를 직접 제어한다. mass floor와 mass trajectory는 all-zero collapse를 막으며, coefficient anchor와 spatial coefficient regularization은 D/r 학습의 식별성을 안정화한다.",
            "checkpoint 수정은 학습 결과 해석에서 가장 중요한 절차 개선이다. 기존 validation MSE 단독 기준은 sparse observation 환경에서 초기 모델을 과대평가할 수 있었다. 새 기준은 validation MSE, late-time physics proxy, PDE residual, IC, front proxy, mass proxy를 log-scale composite score로 합산하고, 최소 epoch 이후의 모델만 checkpoint 후보로 삼는다.",
            "time marching과 time-slab curriculum은 parabolic PDE에서 초반 시간 오차가 후반 front 전체로 전파되는 문제를 줄이기 위한 장치다. 전체 시간 구간을 한 번에 강제하지 않고, 초기장과 쉬운 시간 구간을 먼저 안정화한 뒤 후반 front를 학습한다. front-aware RAR는 residual이 큰 영역뿐 아니라 u(1-u)와 gradient가 큰 active front 주변 collocation을 보강한다.",
            "새로 추가한 discrete RK4 consistency loss는 RK4 결과장을 정답 label로 주입하지 않고, 모델이 예측한 u(t)를 Fisher-KPP PDE의 한 RK4 step으로 전개한 값과 모델의 u(t+dt)를 맞춘다. 이는 discrete-time PINN 계열의 시간 전개 제약을 차용한 것이며, 연속 AD residual만으로 약했던 moving-front phase 누적 오차를 직접 줄이기 위한 공정한 physics loss이다.",
            "산림청 코드에서는 물리 제약을 지도 도메인에 맞게 바꿨다. 바다는 diffusion이 일어날 수 없는 영역으로 처리하고, RK4는 masked no-flux diffusion을 사용하며, PINN은 land-only collocation과 sea exclusion penalty를 사용한다. raw CSV가 있을 때는 조치 시작 전 월별 시간축을 구성해 방제 이후 동역학과 자연 확산 동역학을 섞지 않도록 한다.",
        ],
        (
            ["해결 장치", "코드 반영", "목적"],
            [
                ["Composite checkpoint", "TrainConfig checkpoint weights, train.py score helper", "관측점 MSE 단독 선택 방지"],
                ["Known IC / hard IC", "known_initial_condition_loss, hard_initial_condition", "초기부터 퍼진 haze와 collapse 억제"],
                ["Level-set/profile loss", "front_level_set_alignment_loss, front_profile_alignment_loss", "front phase, level, slope 정렬"],
                ["Support Tversky", "front_support_tversky_loss", "false negative support 누락 감소"],
                ["Mass guards", "mass_floor_trajectory_loss, parabolic_mass_balance_loss", "감염 총량 보존과 all-zero 방지"],
                ["Discrete RK4 consistency", "discrete_rk4_consistency_loss", "RK4 label 없이 시간 전개 일관성 강화"],
                ["RK4 teacher ablation", "rk4_pretrain, rk4_teacher proxy", "default off; 비교 및 ablation 전용"],
                ["Land/sea constraints", "land mask, sea penalty, masked RK4", "한국 지도에서 바다 확산 금지"],
            ],
            [1.75, 3.20, 2.30],
        ),
    ),
    (
        "8. 한국 산림청 데이터와 문제 구성",
        [
            "산림청 계열은 compact GitHub-safe 데이터와 raw CSV 선택 경로를 모두 지원한다. compact 데이터는 2016-2023 감염목 좌표와 연도, manifest, province GeoJSON을 포함한다. raw CSV가 있으면 `조사일자`와 `방제완료여부`를 읽어 조치 전 월별 시간축을 재구성할 수 있다.",
            "공간 좌표는 raw bundle의 EPSG:5181 좌표를 EPSG:5179로 변환하고, 격자 계산에서는 x,y를 [0, 1]로 정규화한다. land mask는 province GeoJSON으로 만들며, observation density는 바다 셀에서 0으로 둔다. RK4는 land/sea 경계에서 masked no-flux diffusion을 쓰고, PINN은 land-only collocation과 sea exclusion penalty를 쓴다.",
            "조치 전 비교는 현재 기본 proxy로 누적 infected and completed records가 50,000을 넘는 첫 날짜를 대규모 조치 시작으로 본다. 제공 raw bundle에서는 이 날짜가 2016-10-02로 추정되며, 기본 비교는 2016-01부터 2016-09까지 월별 관측, RK4, PINN을 같은 elapsed-year grid에서 비교한다.",
        ],
        (
            ["데이터 축", "현재 처리"],
            [
                ["annual compact", "2016-2023 연도별 감염 위치를 GitHub에 포함"],
                ["pre-action month", "raw CSV가 있을 때 조치 시작 전 월별 확산만 비교"],
                ["land/sea", "관측, RK4, PINN 모두 바다 확산을 금지하는 제약 포함"],
            ],
            [1.55, 5.50],
        ),
    ),
    (
        "9. 산림청 PINN과 D/r 해석",
        [
            "산림청 PINN은 `fit_korea_pine_wilt_pinn`에서 동작한다. 모델은 기본 PINN과 같은 `OriginPINN` 계열을 쓰지만, seed-front feature와 travelling-wave exact feature는 끄고, 관측 격자와 land mask에 맞춘 data loss, initial condition loss, land-only PDE residual, sea penalty, weak boundary loss를 사용한다.",
            "계수는 normalized diffusion과 reaction을 학습한다. 물리 해석은 grid의 kilometer scale을 통해 D_phys = D_norm * L_km^2로 되돌린다. x/y 방향의 실제 지도 폭과 높이가 다르기 때문에 normalized PDE residual에서는 D_x = D_phys / width_km^2, D_y = D_phys / height_km^2인 anisotropic scaling을 사용한다.",
            "이 구조는 기존 산림청 zip의 D/r 추정 원리를 일반화한다. 기존 zip이 계수 산정과 forward PINN을 중심으로 구성되었다면, 현재 모델은 land mask, sea penalty, 공간 계수장, monthly action-time comparison, RK4/PINN 공동 시각화를 포함한다. 산림청 결과의 D/r은 관측 밀도, 방제 이력, 신고 체계, 벌목 효과가 결합된 effective spread coefficient로 정리된다.",
        ],
        (
            ["계수", "물리 해석"],
            [
                ["D_norm", "[0,1] 정규화 좌표에서 PINN이 학습하는 기본 diffusion"],
                ["D_phys", "선택한 L_km로 환산한 km^2/year diffusion"],
                ["r", "정규화 시간 단위가 year이므로 reaction_per_year로 해석"],
            ],
            [1.55, 5.50],
        ),
    ),
    (
        "10. 적용 범위와 참고문헌",
        [
            "현재 구현은 Fisher-KPP family를 중심으로 forward field, moving front, 계수 추정, land-constrained spread를 함께 다루는 기준 모델이다. 기본 synthetic 문제에서는 front-aware PINN이 moving front를 추적하고, 산림청 문제에서는 관측 과정과 방제 조치가 반영된 밀도장을 effective reaction-diffusion 계수로 정리한다. 산림청 계열에서 확인된 주요 문제는 바다 확산, 실제 support 누락, 연도 단위 관측의 시간 해상도 부족, 방제 조치 이후 동역학 변화다.",
            "바다 확산은 land mask, masked no-flux RK4, land-only collocation, sea exclusion penalty로 해결한다. support 누락은 false-negative weight가 큰 support Tversky와 mass trajectory로 보정하고, 과도한 false positive는 sea penalty와 support false-positive weight로 제어한다. 연도 단위 관측의 시간 해상도 부족은 raw CSV가 있을 때 pre-action monthly axis로 바꾸어 조치 시작 전 확산만 비교하는 방식으로 정리한다.",
            "확장 방향은 multi-seed full run, 조치 시점 이후의 control term 명시화, 실제 월별 또는 일별 고해상도 관측 확보, front metric 기반 ablation의 통계 반복, D/r posterior와 uncertainty reporting이다. 현재 산출물은 validation figure, GIF, metrics JSON을 통해 field error, front error, mass trajectory, learned physics, checkpoint 선택 근거를 일관된 기준으로 확인한다.",
        ],
        None,
    ),
]

PINN_REFERENCES = [
    "Fisher, Ronald A. 1937. “The Wave of Advance of Advantageous Genes.” Annals of Eugenics 7 (4): 355–369.",
    "Kolmogorov, A. N., I. G. Petrovsky, and N. S. Piskunov. 1937. “A Study of the Equation of Diffusion with Increase in the Quantity of Matter, and Its Application to a Biological Problem.” Bulletin of Moscow University, Mathematics and Mechanics 1: 1–25.",
    "Ablowitz, Mark J., and Anthony Zeppetella. 1979. “Explicit Solutions of Fisher’s Equation for a Special Wave Speed.” Bulletin of Mathematical Biology 41: 835–840.",
    "Raissi, Maziar, Paris Perdikaris, and George Em Karniadakis. 2019. “Physics-Informed Neural Networks: A Deep Learning Framework for Solving Forward and Inverse Problems Involving Nonlinear Partial Differential Equations.” Journal of Computational Physics 378: 686–707.",
    "Wang, Sifan, Yujun Teng, and Paris Perdikaris. 2021. “Understanding and Mitigating Gradient Flow Pathologies in Physics-Informed Neural Networks.” SIAM Journal on Scientific Computing 43 (5): A3055–A3081.",
    "Krishnapriyan, Aditi S., Amir Gholami, Shandian Zhe, Robert M. Kirby, and Michael W. Mahoney. 2021. “Characterizing Possible Failure Modes in Physics-Informed Neural Networks.” Advances in Neural Information Processing Systems 34.",
    "Jagtap, Ameya D., and George Em Karniadakis. 2020. “Extended Physics-Informed Neural Networks (XPINNs): A Generalized Space-Time Domain Decomposition Based Deep Learning Framework for Nonlinear Partial Differential Equations.” Communications in Computational Physics 28 (5): 2002–2041.",
    "Moseley, Ben, Andrew Markham, and Tarje Nissen-Meyer. 2023. “Finite Basis Physics-Informed Neural Networks (FBPINNs): A Scalable Domain Decomposition Approach for Solving Differential Equations.” Advances in Computational Mathematics 49: 62.",
    "Yu, Jeremy, Lu Lu, Xuhui Meng, and George Em Karniadakis. 2022. “Gradient-Enhanced Physics-Informed Neural Networks for Forward and Inverse PDE Problems.” Computer Methods in Applied Mechanics and Engineering 393: 114823.",
    "Wu, Chenxi, Min Zhu, Qinyang Tan, Yadhu Kartha, and Lu Lu. 2023. “A Comprehensive Study of Non-Adaptive and Residual-Based Adaptive Sampling for Physics-Informed Neural Networks.” Computer Methods in Applied Mechanics and Engineering 403: 115671.",
    "Wang, Sifan, Shyam Sankaran, and Paris Perdikaris. 2024. “Respecting Causality for Training Physics-Informed Neural Networks.” Computer Methods in Applied Mechanics and Engineering 421: 116813.",
    "Akrivis, Georgios, Charalambos G. Makridakis, and Costas Smaragdakis. 2025. “Runge-Kutta Physics-Informed Neural Networks: Formulation and Analysis.” Numerische Mathematik.",
    "Wang, Sifan, Xinling Yu, and Paris Perdikaris. 2024. “PirateNets: Physics-Informed Deep Learning with Residual Adaptive Networks.” Journal of Machine Learning Research 25.",
]


RK4_PAGES = [
    (
        "1. Fisher-KPP 검증 regime",
        [
            "본 RK4 모듈의 기본 검증 문제는 Fisher-KPP 반응-확산 방정식 u_t = D Δu + r u(1-u)이다. 이번 수정에서는 시간 적분법을 변경하지 않고, 검증 regime만 폐형식 travelling-wave 해가 존재하는 문제로 정렬하였다. 목적은 수치해가 정확해와 직접 비교될 수 있도록 초기조건, Dirichlet boundary condition, 최종 시간, 격자, 오차 지표를 모두 같은 해에서 유도하는 것이다.",
            "1차원 기준 문제는 Ablowitz-Zeppetella travelling wave이다. 확산계수와 반응계수는 D = 1, r = 1로 두며, 특수 파속 c = 5/sqrt(6)에서 u(x,t) = {1 + exp[(x - c t - x0)/sqrt(6)]}^{-2}가 정확해가 된다. 계산 영역은 x in [-20, 20], 최종 시간은 T = 10, 기본 격자는 Nx = 201, 시간 간격은 dt = 0.005이다.",
            "2차원 기준 문제는 generalized Fisher-KPP exact wave의 p = 1 경우이다. 방정식은 u_t = u_xx + u_yy + u(1-u)이고, phi = pi/4, C = 0을 두면 u(x,y,t) = [0.5 tanh((x+y)/(4 sqrt(3)) + 5t/12) + 0.5]^2가 정확해이다. 계산 영역은 x,y in [-15, 15], 최종 시간은 T = 3, 기본 격자는 61 x 61, 시간 간격은 dt = 0.01이다.",
            "두 문제 모두 초기조건은 정확해의 t = 0 단면을 사용한다. 1D에서는 양 끝점 x = -L/2, x = L/2의 Dirichlet 값을 정확해에서 가져오고, 2D에서는 네 변 x = -L/2, x = L/2, y = -L/2, y = L/2의 Dirichlet 값을 같은 2D 정확해에서 가져온다. 따라서 경계 오차와 내부 시간 적분 오차를 분리하여 평가할 수 있다.",
        ],
        (
            ["항목", "1D 기준", "2D 기준"],
            [
                ["PDE", "u_t = u_xx + u(1-u)", "u_t = u_xx + u_yy + u(1-u)"],
                ["Exact solution", "Ablowitz-Zeppetella exponential wave", "Generalized Fisher-KPP tanh wave, p=1"],
                ["Domain", "[-20, 20]", "[-15, 15] x [-15, 15]"],
                ["Default grid/time", "Nx=201, T=10, dt=0.005", "61x61, T=3, dt=0.01"],
            ],
            [1.35, 2.95, 3.25],
        ),
    ),
    (
        "2. 공간 이산화와 Dirichlet boundary",
        [
            "공간 이산화는 method-of-lines 관점에서 수행된다. 1D 격자는 x_i = -L/2 + i Δx, i = 0, 1, ..., Nx - 1로 정의하고, 내부점 i = 1, ..., Nx - 2에서 u_xx ≈ (u_{i-1} - 2u_i + u_{i+1}) / Δx^2의 중심차분을 사용한다. 양 끝점은 미지수가 아니라 정확해에서 산출된 boundary value로 고정된다.",
            "1D Dirichlet boundary contribution을 포함하면 내부점에 대한 Laplacian matrix는 삼중대각 행렬이 된다. 본 RK4 구현에서는 전체 벡터에 boundary 값을 먼저 삽입하고, 내부점에 대해서만 RHS를 계산한다. 이 방식은 행렬을 명시적으로 구성하지 않아도 동일한 중심차분 stencil과 boundary contribution을 재현한다.",
            "2D에서는 동일한 중심차분 stencil을 x, y 방향에 적용한다. 내부 격자점에서 Δu ≈ (u_{i-1,j} - 2u_{i,j} + u_{i+1,j})/Δx^2 + (u_{i,j-1} - 2u_{i,j} + u_{i,j+1})/Δy^2이며, 정사각 균일 격자를 사용하므로 Δx = Δy이다. 행렬 관점에서는 L_2D = L_x ⊗ I_y + I_x ⊗ L_y로 해석된다.",
            "이번 regime에서 2D boundary는 no-flux가 아니라 exact Dirichlet boundary이다. 매 RK4 stage마다 stage time에 해당하는 정확해를 네 변에 다시 삽입한다. 따라서 수치해가 받는 경계 정보는 분석해와 일치하며, final relative L2와 absolute error map은 공간 이산화와 RK4 시간 적분에서 발생한 오차를 주로 반영한다.",
        ],
        (
            ["구분", "적용 방식"],
            [
                ["1D 내부점", "second-order central finite difference"],
                ["1D boundary", "u(-L/2,t), u(L/2,t)를 정확해에서 부여"],
                ["2D 내부점", "L_x ⊗ I_y + I_x ⊗ L_y 형태의 five-point Laplacian"],
                ["2D boundary", "네 변 모두 generalized Fisher-KPP 정확해에서 부여"],
            ],
            [2.10, 4.95],
        ),
    ),
    (
        "3. RK4 시간 적분과 안정성 판정",
        [
            "시간 적분은 고전적 4단계 Runge-Kutta 방법을 그대로 사용한다. 반이산화된 ODE를 u_t = F(u,t)로 쓰면 k1 = F(u^n,t_n), k2 = F(u^n + dt k1/2, t_n + dt/2), k3 = F(u^n + dt k2/2, t_n + dt/2), k4 = F(u^n + dt k3, t_n + dt)이고, u^{n+1} = u^n + dt(k1 + 2k2 + 2k3 + k4)/6으로 갱신한다.",
            "Dirichlet 문제에서는 각 stage 입력장에 해당 stage time의 boundary value를 먼저 삽입한 뒤 RHS를 계산한다. 따라서 k1, k2, k3, k4가 모두 같은 물리적 boundary condition 아래에서 계산된다. 이 절차는 방법을 바꾸는 것이 아니라, 정확해 기반 검증 문제에 맞도록 stage boundary를 일관되게 적용하는 것이다.",
            "명시적 RK4는 diffusion-dominated semi-discrete system에서 안정성 제약을 가진다. 코드의 실용 판정은 dt_diff_limit = 0.69 Δx^2 / (dim D)와 reaction scale 1/r 중 작은 값을 safety factor와 비교한다. 1D 기본값 dt = 0.005와 2D 기본값 dt = 0.01은 해당 기준 안에 있으며, 안정성 위반 시 실행 전에 오류를 발생시킨다.",
            "오차 평가는 exact final field와 numerical final field의 relative L2, final absolute error, 1D front position, 2D mass 및 front-area 지표로 수행한다. 정확해 기반 regime이므로 RK4 결과는 단순한 reference curve가 아니라 분석해와 직접 대조되는 수치해이다. 별도의 forward Euler, backward Euler, trapezoidal 비교에서는 같은 grid, 같은 dt, 같은 초기조건을 유지하여 시간 적분법의 차이만 관찰한다.",
        ],
        (
            ["항목", "검증 원칙"],
            [
                ["Method", "classical RK4, four-stage explicit time integration"],
                ["Stage boundary", "각 stage time의 exact Dirichlet 값을 삽입"],
                ["Stability", "diffusion scale과 reaction scale을 동시에 점검"],
                ["Metrics", "relative L2, absolute error, front/mass diagnostics"],
            ],
            [1.75, 5.30],
        ),
    ),
    (
        "4. 결과 해석과 보고 기준",
        [
            "1D 결과에서 front가 오른쪽으로 이동하고 final relative L2가 작은 값으로 유지되면, 중심차분 공간 이산화와 RK4 시간 적분이 Ablowitz-Zeppetella travelling wave의 phase와 profile을 함께 포착하고 있음을 의미한다. 격자를 세분화하면 중심차분의 2차 공간 정확도에 따라 오차가 감소하는 경향을 보이며, dt를 충분히 줄인 조건에서는 공간 오차가 전체 오차를 지배한다.",
            "2D 결과에서 tanh front가 대각선 방향으로 이동하고 exact final field와 수치해의 차이가 작으면, Kronecker 형태의 2D Laplacian stencil과 stage별 exact Dirichlet 처리가 정상적으로 작동한다고 해석한다. 2D 문제는 내부 격자점 수가 (Nx - 2)(Ny - 2)에 비례하므로 격자 세분화는 정확도 향상과 계산비 증가를 동시에 유발한다.",
            "시간 수렴 차수는 exact solution 대비 오차가 아니라 같은 공간 격자에서 매우 작은 시간간격으로 계산한 reference solution과의 final-time relative L2 차이로 측정한다. 1D에서는 Nx = 201, dt_ref = 0.00125를 기준으로 dt = 0.02, 0.01, 0.005, 0.0025 해를 비교하고, 2D에서는 Nx = Ny = 121, dt_ref = 0.0025를 기준으로 dt = 0.02, 0.01, 0.005 해를 비교한다. 관측 시간 수렴차수는 p_i = log(E_i/E_{i+1}) / log(dt_i/dt_{i+1})로 계산한다. 현재 실행 결과 1D는 약 4.57, 4.25, 4.21, 2D는 약 4.34, 4.24의 차수를 보여 classical RK4의 4차 시간 정확도와 일치한다.",
            "보고서와 notebook의 핵심 표기는 PDE, 정확해, 초기조건, boundary condition, 공간 이산화, RK4 stage 적용 순서, 안정성 판정, 오차 지표의 순서로 구성한다. 이는 사진의 설명 흐름을 따르되, 연구 재현성을 위해 모든 상수와 비교 기준을 명시한 형식이다. 본 문서의 수식과 코드 설정은 fisher-kpp-rk4 패키지의 현재 기본값과 일치한다.",
            "이 기준은 RK4를 다른 방법으로 대체하지 않는다. 변경된 것은 검증 regime이며, solver는 여전히 method-of-lines, second-order central finite difference, classical RK4를 사용한다. 따라서 기존 RK4 결과와 새 결과의 차이는 수치법 변경이 아니라 정확해를 갖는 Fisher-KPP benchmark로 문제 설정을 정렬한 데서 발생한다.",
        ],
        None,
    ),
]

RK4_REFERENCES = [
    "Fisher, Ronald A. 1937. “The Wave of Advance of Advantageous Genes.” Annals of Eugenics 7 (4): 355–369.",
    "Kolmogorov, A. N., I. G. Petrovsky, and N. S. Piskunov. 1937. “A Study of the Equation of Diffusion with Increase in the Quantity of Matter, and Its Application to a Biological Problem.” Bulletin of Moscow University, Mathematics and Mechanics 1: 1–25.",
    "Ablowitz, Mark J., and Anthony Zeppetella. 1979. “Explicit Solutions of Fisher’s Equation for a Special Wave Speed.” Bulletin of Mathematical Biology 41: 835–840.",
    "Butcher, John C. 2016. Numerical Methods for Ordinary Differential Equations. 3rd ed. Chichester: Wiley.",
    "Morton, K. W., and D. F. Mayers. 2005. Numerical Solution of Partial Differential Equations: An Introduction. 2nd ed. Cambridge: Cambridge University Press.",
    "Zhou, Hanyu, and Yu-Xin Ren. 2024. “High-Order Accurate Implicit Scheme Based on Temporal Reconstruction for Solving Compressible Navier-Stokes Equations.” In Proceedings of the 12th International Conference on Computational Fluid Dynamics, Kobe, Japan.",
]


def _build_pages(document: Document, pages, references: list[str] | None = None) -> None:
    for idx, (heading, paragraphs, table) in enumerate(pages):
        _add_heading(document, heading)
        for paragraph in paragraphs:
            _add_para(document, paragraph)
        if table is not None:
            headers, rows, widths = table
            _add_table(document, headers, rows, widths)
        if references and idx == len(pages) - 1:
            _add_heading(document, "참고문헌", level=2)
            for reference in references:
                _add_reference(document, reference)
        if idx != len(pages) - 1:
            document.add_page_break()


def _audit_all_runs(document: Document) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _format_run(run, bold=bool(run.bold), italic=bool(run.italic))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _format_run(run, bold=bool(run.bold), italic=bool(run.italic))


def build_pinn_doc(path: Path) -> None:
    document = Document()
    _configure_document(document, "기본 PINN 및 산림청 PINN 구현 설명")
    _add_title(document, "기본 PINN 및 산림청 PINN 구현 설명", "Fisher-KPP 및 Korea pine-wilt PINN 구현 사양")
    _build_pages(document, PINN_PAGES, PINN_REFERENCES)
    _audit_all_runs(document)
    document.save(path)


def build_rk4_doc(path: Path) -> None:
    document = Document()
    _configure_document(document, "Fisher-KPP RK4 검증 regime 및 수치해석 사양")
    _add_title(document, "Fisher-KPP RK4 검증 regime 및 수치해석 사양", "Exact travelling-wave benchmark 기반 1D/2D RK4 보고서")
    _build_pages(document, RK4_PAGES, RK4_REFERENCES)
    _audit_all_runs(document)
    document.save(path)


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    build_pinn_doc(DOCS_DIR / "fisher_kpp_pinn_and_korea_pinn_technical_note.docx")
    build_rk4_doc(DOCS_DIR / "fisher_kpp_rk4_pde_technical_note.docx")


if __name__ == "__main__":
    main()

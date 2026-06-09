from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT_NAME = "Malgun Gothic"
FONT_SIZE = 10.0
BLACK = RGBColor(0, 0, 0)


def _set_run_font(run, *, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(FONT_SIZE)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold


def _set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 4, line: float = 1.08) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = document.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(FONT_SIZE)
        style.font.color.rgb = BLACK
        style.paragraph_format.line_spacing = 1.08
        style.paragraph_format.space_after = Pt(4)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.bold = True

    styles["Heading 1"].paragraph_format.space_before = Pt(9)
    styles["Heading 1"].paragraph_format.space_after = Pt(4)
    styles["Heading 2"].paragraph_format.space_before = Pt(7)
    styles["Heading 2"].paragraph_format.space_after = Pt(3)
    styles["Heading 3"].paragraph_format.space_before = Pt(5)
    styles["Heading 3"].paragraph_format.space_after = Pt(2)


def _write_inline_markdown(paragraph, text: str, *, bold: bool = False) -> None:
    parts = text.split("`")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        _set_run_font(run, bold=bold)
        if idx % 2 == 1:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(FONT_SIZE)
            run.font.color.rgb = BLACK


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(paragraph, before=0, after=8, line=1.08)
    run = paragraph.add_run(text)
    _set_run_font(run, bold=True)


def build_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    _configure_document(document)
    lines = markdown_path.read_text(encoding="utf-8").splitlines()

    in_code = False
    code_buffer: list[str] = []
    title_done = False

    def flush_code() -> None:
        nonlocal code_buffer
        if not code_buffer:
            return
        paragraph = document.add_paragraph()
        _set_paragraph_spacing(paragraph, before=2, after=5, line=1.0)
        run = paragraph.add_run("\n".join(code_buffer))
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(FONT_SIZE)
        run.font.color.rgb = BLACK
        code_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
                code_buffer = []
            continue
        if in_code:
            code_buffer.append(line)
            continue
        if not line.strip():
            continue

        if line.startswith("# "):
            _add_title(document, line[2:].strip())
            title_done = True
            continue

        if line.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            _write_inline_markdown(paragraph, line[4:].strip(), bold=True)
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            _write_inline_markdown(paragraph, line[3:].strip(), bold=True)
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0.16 if title_done else 0)
        _set_paragraph_spacing(paragraph, after=4, line=1.08)
        _write_inline_markdown(paragraph, line)

    flush_code()

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    run = footer.add_run("Fisher-KPP PINN 및 한국 산림청 확장 모델 기술 설명서")
    _set_run_font(run)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the Fisher-KPP PINN technical explanation DOCX.")
    parser.add_argument("--input", type=Path, default=Path("docs") / "fisher_kpp_pinn_review_response.md")
    parser.add_argument("--output", type=Path, default=Path("docs") / "fisher_kpp_pinn_review_response.docx")
    args = parser.parse_args()
    build_docx(args.input, args.output)
    print(args.output)


if __name__ == "__main__":
    main()
